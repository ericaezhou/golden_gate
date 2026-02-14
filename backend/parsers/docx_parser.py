"""DOCX parser for extracting structured content from Word documents."""

import os
import re
from typing import List
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from backend.parsers import register_parser, ParseResult


def extract_paragraph_text(paragraph: Paragraph) -> str:
    """Extract text from a paragraph, preserving bold formatting."""
    result = []
    for run in paragraph.runs:
        text = run.text
        if run.bold:
            result.append(f"**{text}**")
        else:
            result.append(text)
    return "".join(result)


def table_to_markdown(table: Table) -> str:
    """Convert a DOCX table to markdown format."""
    if not table.rows:
        return ""

    lines = []

    # Process all rows
    for row_idx, row in enumerate(table.rows):
        cells = []
        for cell in row.cells:
            # Extract text from all paragraphs in the cell
            cell_text = " ".join(p.text.strip() for p in cell.paragraphs).strip()
            cells.append(cell_text)

        # Add the row
        lines.append("| " + " | ".join(cells) + " |")

        # Add separator after first row (header)
        if row_idx == 0:
            lines.append("| " + " | ".join(["---"] * len(cells)) + " |")

    return "\n".join(lines)


def paragraph_to_markdown(paragraph: Paragraph) -> str:
    """Convert a paragraph to markdown, handling heading styles."""
    text = extract_paragraph_text(paragraph)

    if not text.strip():
        return ""

    # Map heading styles to markdown
    style_name = paragraph.style.name if paragraph.style else ""

    if style_name == "Heading 1":
        return f"# {text}"
    elif style_name == "Heading 2":
        return f"## {text}"
    elif style_name == "Heading 3":
        return f"### {text}"
    else:
        return text


def _extract_file_references(text: str) -> list[str]:
    """Extract unique filename references from text."""
    pattern = re.compile(
        r'\b[\w.-]+\.(xlsx|xls|csv|py|sql|ipynb|pdf|pptx|docx|txt|db)\b',
        re.IGNORECASE,
    )
    seen = set()
    result = []
    for m in pattern.finditer(text):
        ref = m.group(0)
        if ref.lower() not in seen:
            seen.add(ref.lower())
            result.append(ref)
    return result


@register_parser("docx")
def parse_docx(path: str) -> ParseResult:
    """
    Parse a DOCX file and extract structured content.

    Args:
        path: Path to the DOCX file

    Returns:
        ParseResult containing metadata, content (as markdown), and references
    """
    doc = Document(path)

    # Initialize metadata counters
    paragraph_count = 0
    table_count = 0
    headings = []

    # Content parts (will be joined into markdown)
    content_parts = []

    # All text for reference extraction
    all_text = []

    # Iterate through document body elements in order
    for element in doc.element.body:
        if isinstance(element, CT_P):
            # It's a paragraph
            paragraph = Paragraph(element, doc)
            para_text = extract_paragraph_text(paragraph)

            if para_text.strip():
                paragraph_count += 1
                all_text.append(para_text)

                # Check if it's a heading
                style_name = paragraph.style.name if paragraph.style else ""
                if style_name in ["Heading 1", "Heading 2", "Heading 3"]:
                    headings.append(para_text)

                # Convert to markdown
                md_text = paragraph_to_markdown(paragraph)
                if md_text:
                    content_parts.append(md_text)

        elif isinstance(element, CT_Tbl):
            # It's a table
            table = Table(element, doc)
            table_count += 1

            # Extract text from table for references
            for row in table.rows:
                for cell in row.cells:
                    cell_text = " ".join(p.text.strip() for p in cell.paragraphs)
                    if cell_text.strip():
                        all_text.append(cell_text)

            # Convert to markdown
            md_table = table_to_markdown(table)
            if md_table:
                content_parts.append(md_table)

    # Join all content with blank lines between elements
    content = "\n\n".join(content_parts)

    # Extract references from all text
    all_text_combined = "\n".join(all_text)
    references = _extract_file_references(all_text_combined)

    # Build metadata
    metadata = {
        "paragraph_count": paragraph_count,
        "table_count": table_count,
        "headings": headings
    }

    # Get file info
    file_id = os.path.basename(path)
    file_path = os.path.abspath(path)

    return ParseResult(
        file_id=file_id,
        file_path=file_path,
        file_type="docx",
        metadata=metadata,
        content=content,
        references=references
    )
